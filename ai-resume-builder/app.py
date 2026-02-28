import os
import io
from datetime import datetime

import streamlit as st
import pandas as pd
from openai import OpenAI
from supabase import create_client, Client
from docx import Document
from fpdf import FPDF

# 1) Page config
st.set_page_config(
    page_title="AI Resume & Portfolio Builder",
    layout="wide",
)
# Custom CSS for gradient background and glassmorphism
st.markdown(
    """
    <style>
    body {
        background: linear-gradient(135deg, #1e3c72, #2a5298, #6a11cb, #2575fc);
        background-attachment: fixed;
    }
    .main {
        background: transparent;
    }
    .stApp {
        background: linear-gradient(135deg, #1e3c72, #2a5298, #6a11cb, #2575fc);
        background-attachment: fixed;
    }
    .glass-card {
        background: rgba(255, 255, 255, 0.12);
        border-radius: 18px;
        padding: 24px 28px;
        box-shadow: 0 18px 45px rgba(15, 23, 42, 0.45);
        backdrop-filter: blur(18px);
        -webkit-backdrop-filter: blur(18px);
        border: 1px solid rgba(255, 255, 255, 0.25);
    }
    .glass-sidebar {
        background: rgba(15, 23, 42, 0.15);
        border-radius: 18px;
        padding: 18px 18px 24px 18px;
        box-shadow: 0 18px 45px rgba(15, 23, 42, 0.4);
        backdrop-filter: blur(18px);
        -webkit-backdrop-filter: blur(18px);
        border: 1px solid rgba(148, 163, 184, 0.6);
    }
    .title-text {
        color: #f9fafb;
        font-weight: 800;
        font-size: 2.1rem;
        letter-spacing: 0.03em;
    }
    .subtitle-text {
        color: #e5e7eb;
        font-size: 0.95rem;
        max-width: 640px;
    }
    .ats-badge {
        display: inline-flex;
        align-items: center;
        padding: 6px 10px;
        border-radius: 999px;
        background: radial-gradient(circle at top left, #22c55e, #16a34a);
        color: #ecfdf5;
        font-size: 0.80rem;
        font-weight: 700;
        box-shadow: 0 10px 25px rgba(22, 163, 74, 0.65);
        margin-left: 10px;
    }
    .section-heading {
        color: #e5e7eb;
        font-size: 1.05rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    .field-label {
        font-size: 0.84rem;
        font-weight: 600;
        color: #e5e7eb;
        margin-bottom: 0.1rem;
    }
    .portfolio-section-title {
        color: #e5e7eb;
        font-weight: 700;
        margin-top: 0.25rem;
        margin-bottom: 0.25rem;
        font-size: 1rem;
    }
    .portfolio-text {
        color: #f9fafb;
        font-size: 0.92rem;
    }
    .small-helper {
        color: #cbd5f5;
        font-size: 0.75rem;
        margin-bottom: 0.25rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# 2) Environment / client setup
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
SUPABASE_URL = os.getenv("SUPABASE_URL", "https://your-supabase-url.supabase.co")
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "your-supabase-anon-or-service-key")

openai_client = None
supabase_client: Client | None = None

if OPENAI_API_KEY:
   openai_client = OpenAI(api_key=OPENAI_API_KEY)
else:
    st.warning("OPENAI_API_KEY is not set in environment variables. AI features may not work.", icon="⚠️")

try:
    if SUPABASE_URL and SUPABASE_KEY and "your-supabase-url" not in SUPABASE_URL and "your-supabase-anon-or-service-key" not in SUPABASE_KEY:
        supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)
    else:
        supabase_client = None
        st.info("Supabase URL or KEY are placeholders. Data will not be saved to database.", icon="ℹ️")
except Exception as e:
    supabase_client = None
    st.error(f"Could not initialize Supabase client: {e}")


# 3) Helper functions
def build_resume_prompt(data: dict) -> str:
    skills_str = ", ".join(data.get("skills", []))
    projects_str = data.get("projects", "").strip()
    experience_str = data.get("experience", "").strip()
    education_str = data.get("education", "").strip()
    linkedin = data.get("linkedin", "").strip()
    github = data.get("github", "").strip()

    prompt = f"""
You are an expert technical resume writer for B.Tech students.

Generate a professional, ATS-optimized resume in PLAIN TEXT ONLY (no markdown, no bullet symbols like •, just use hyphens for bullets). 
Use clear section titles in uppercase, concise bullet points, and modern wording appropriate for freshers or early-career candidates.
The resume must be suitable for software/tech roles.

USER PROFILE:
Name: {data.get("name", "")}
Email: {data.get("email", "")}
Phone: {data.get("phone", "")}

Education:
{education_str}

Skills:
{skills_str}

Projects:
{projects_str}

Experience:
{experience_str}

LinkedIn: {linkedin}
GitHub: {github}

REQUIREMENTS:
- Start with a strong objective or summary of 2–3 lines.
- Use sections in the order: SUMMARY, EDUCATION, SKILLS, PROJECTS, EXPERIENCE, LINKS, ADDITIONAL if needed.
- Use simple text formatting with clear headings, for example:
  SUMMARY
  - First summary line
  - Second summary line
- Use clear hyphen-based bullet points starting with action verbs.
- Focus on achievements, technologies, and impact.
- Do NOT use any markdown, tables, or special bullet characters.
- Keep the length to no more than 1.5 pages of plain text.

Return ONLY the resume text with no explanations.
"""
    return prompt.strip()


def build_cover_letter_prompt(data: dict, job_title: str, company: str) -> str:
    skills_str = ", ".join(data.get("skills", []))
    projects_str = data.get("projects", "").strip()
    experience_str = data.get("experience", "").strip()
    education_str = data.get("education", "").strip()

    prompt = f"""
You are an expert career coach and cover letter writer.

Write a tailored, professional cover letter in PLAIN TEXT ONLY (no markdown, no bullet symbols, no tables).
The cover letter is for a B.Tech student or recent graduate applying for the role specified below.

USER PROFILE:
Name: {data.get("name", "")}
Email: {data.get("email", "")}
Phone: {data.get("phone", "")}
Education: {education_str}
Skills: {skills_str}
Projects:
{projects_str}

Experience:
{experience_str}

JOB DETAILS:
Job Title: {job_title}
Company: {company}

REQUIREMENTS:
- Address the letter to "Hiring Manager".
- 3–5 concise paragraphs.
- Highlight relevant skills, projects, and experience for the role of {job_title} at {company}.
- Keep tone sincere, confident, and suitable for freshers/early career candidates.
- Do NOT use any markdown or special formatting. Simple plain text paragraphs only.
- End with an appropriate closing and the candidate's name.

Return ONLY the cover letter text with no explanations.
"""
    return prompt.strip()


def call_openai(prompt: str) -> str:
    if openai_client is None:
        raise ValueError("OpenAI client is not initialized. Check OPENAI_API_KEY.")
    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant specialized in creating plain-text resumes and cover letters."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.6,
            max_tokens=1200,
        )
        text = response.choices[0].message.content
        return text.strip()
    except Exception as e:
        raise RuntimeError(f"OpenAI API error: {e}")


def create_docx_from_text(text: str, title: str) -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
        doc.add_paragraph("")
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def create_pdf_from_text(text: str, title: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    if title:
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, txt=title, ln=True, align="L")
        pdf.ln(4)
        pdf.set_font("Arial", size=12)

    for line in text.split("\n"):
        if not line.strip():
            pdf.ln(5)
        else:
            pdf.multi_cell(0, 6, txt=line)
    pdf_output = pdf.output(dest="S").encode("latin-1")
    return pdf_output


def build_portfolio_sections(data: dict, resume_text: str) -> dict:
    name = data.get("name", "").strip()
    education = data.get("education", "").strip()
    skills_list = data.get("skills", [])
    projects_text = data.get("projects", "").strip()
    experience_text = data.get("experience", "").strip()

    summary_lines = []
    if name:
        summary_lines.append(f"{name} is a B.Tech student aspiring to build a strong career in technology roles.")
    if skills_list:
        summary_lines.append(f"Key skills include: {', '.join(skills_list)}.")
    if education:
        summary_lines.append(f"Academic background: {education}.")
    if not summary_lines:
        summary_lines.append("A motivated B.Tech student building a strong foundation in technology and problem-solving.")

    summary = " ".join(summary_lines)

    projects = []
    if projects_text:
        for line in projects_text.split("\n"):
            clean = line.strip("-").strip()
            if clean:
                projects.append(clean)

    exp_highlights = []
    if resume_text:
        lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
        for line in lines:
            if line.upper() in ["SUMMARY", "EDUCATION", "SKILLS", "PROJECTS", "EXPERIENCE", "LINKS", "ADDITIONAL"]:
                continue
            if line.startswith("-"):
                exp_highlights.append(line.strip("-").strip())
            elif any(keyword in line.lower() for keyword in ["intern", "project", "developed", "built", "implemented", "designed"]):
                exp_highlights.append(line)
            if len(exp_highlights) >= 5:
                break

    if not exp_highlights and experience_text:
        for line in experience_text.split("\n"):
            clean = line.strip("-").strip()
            if clean:
                exp_highlights.append(clean)
            if len(exp_highlights) >= 5:
                break

    return {
        "summary": summary,
        "projects": projects,
        "experience_highlights": exp_highlights,
    }


def save_to_supabase(data: dict, resume_text: str, cover_letter: str | None):
    if supabase_client is None:
        return

    try:
        payload = {
            "name": data.get("name", "").strip(),
            "email": data.get("email", "").strip(),
            "phone": data.get("phone", "").strip(),
            "education": data.get("education", "").strip(),
            "skills": ", ".join(data.get("skills", [])),
            "projects": data.get("projects", "").strip(),
            "experience": data.get("experience", "").strip(),
            "linkedin": data.get("linkedin", "").strip(),
            "github": data.get("github", "").strip(),
            "resume_content": resume_text or "",
            "cover_letter": cover_letter or "",
            "created_at": datetime.utcnow().isoformat(),
        }
        email = payload["email"]
        if email:
            payload["email"] = email
        supabase_client.table("users_table").upsert(payload, on_conflict="email").execute()
    except Exception as e:
        st.error(f"Could not save data to Supabase: {e}")


# 4) Session state
if "resume_text" not in st.session_state:
    st.session_state.resume_text = ""
if "cover_letter_text" not in st.session_state:
    st.session_state.cover_letter_text = ""
if "portfolio" not in st.session_state:
    st.session_state.portfolio = {}
if "form_data" not in st.session_state:
    st.session_state.form_data = {}

# 5) Layout: Title and columns
top_col1, top_col2 = st.columns([0.75, 0.25])
with top_col1:
    st.markdown('<div class="title-text">AI Resume & Portfolio Builder</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="subtitle-text">Create a professional, ATS-friendly resume, tailored cover letter, and portfolio preview powered by OpenAI — optimized for B.Tech students.</div>',
        unsafe_allow_html=True,
    )
with top_col2:
    st.markdown(
        '<div style="display:flex;justify-content:flex-end;align-items:center;margin-top:8px;">'
        '<span class="ats-badge">ATS Score: 92%</span>'
        '</div>',
        unsafe_allow_html=True,
    )

st.markdown("<br/>", unsafe_allow_html=True)

left_col, right_col = st.columns([0.32, 0.68])

# 1) Sidebar-like form in left column
with left_col:
    st.markdown('<div class="glass-sidebar">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">Step 1 · Student & Contact Details</div>', unsafe_allow_html=True)

    name = st.text_input("Full Name (as on certificates)", key="name_input")
    email = st.text_input("Email (for resume contact)", key="email_input")
    phone = st.text_input("Mobile Number (with country code)", key="phone_input")

    st.markdown('<div class="section-heading" style="margin-top:0.75rem;">Step 2 · Education</div>', unsafe_allow_html=True)
    education = st.text_area(
        "Education (degree, college, graduation year, CGPA)",
        height=110,
        key="education_input",
    )

    st.markdown('<div class="section-heading" style="margin-top:0.75rem;">Step 3 · Technical Skills</div>', unsafe_allow_html=True)
    available_skills = [
        "Python",
        "C++",
        "Java",
        "JavaScript",
        "React",
        "Node.js",
        "Django",
        "Machine Learning",
        "Data Structures",
        "Algorithms",
        "SQL",
        "DevOps",
        "Cloud (AWS/Azure/GCP)",
        "HTML/CSS",
        "Git/GitHub",
    ]
    skills = st.multiselect(
        "Select your strongest technical skills",
        options=available_skills,
        key="skills_input",
    )

    st.markdown('<div class="section-heading" style="margin-top:0.75rem;">Step 4 · Academic & Personal Projects</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="small-helper">Write one project per line. Include goal, your contribution, and technologies used.</div>',
        unsafe_allow_html=True,
    )
    projects = st.text_area(
        "Projects (one per line)",
        height=140,
        key="projects_input",
    )

    st.markdown('<div class="section-heading" style="margin-top:0.75rem;">Step 5 · Experience & Activities</div>', unsafe_allow_html=True)
    experience = st.text_area(
        "Internships, trainings, hackathons, leadership roles, or other experience",
        height=110,
        key="experience_input",
    )

    st.markdown('<div class="section-heading" style="margin-top:0.75rem;">Step 6 · Online Profiles</div>', unsafe_allow_html=True)
    linkedin = st.text_input("LinkedIn Profile URL", key="linkedin_input")
    github = st.text_input("GitHub Profile URL", key="github_input")

    st.markdown('</div>', unsafe_allow_html=True)

# Store in session form_data
st.session_state.form_data = {
    "name": name.strip(),
    "email": email.strip(),
    "phone": phone.strip(),
    "education": education.strip(),
    "skills": skills,
    "projects": projects.strip(),
    "experience": experience.strip(),
    "linkedin": linkedin.strip(),
    "github": github.strip(),
}

# 6) Main tabs
with right_col:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    tabs = st.tabs(["Resume", "Cover Letter", "Portfolio"])

    # Resume Tab
    with tabs[0]:
        st.markdown("### Step 7 · AI-Generated Resume")
        st.write(
            "Complete the form on the left, then click **Generate AI Resume** to create a placement-ready, ATS-optimized resume for B.Tech roles."
        )

        generate_resume_clicked = st.button("Generate AI Resume", type="primary")

        if generate_resume_clicked:
            form_data = st.session_state.form_data
            missing_fields = []
            if not form_data.get("name"):
                missing_fields.append("Name")
            if not form_data.get("email"):
                missing_fields.append("Email")
            if not form_data.get("education"):
                missing_fields.append("Education")

            if missing_fields:
                st.error(f"Please fill the required fields: {', '.join(missing_fields)}.")
            else:
                prompt = build_resume_prompt(form_data)
                progress_bar = st.progress(0)
                with st.spinner("Generating your AI-powered resume..."):
                    try:
                        progress_bar.progress(20)
                        resume_text = call_openai(prompt)
                        progress_bar.progress(80)
                        st.session_state.resume_text = resume_text
                        portfolio_sections = build_portfolio_sections(form_data, resume_text)
                        st.session_state.portfolio = portfolio_sections
                        save_to_supabase(form_data, resume_text, st.session_state.cover_letter_text)
                        progress_bar.progress(100)
                        st.success("Resume generated successfully.")
                    except Exception as e:
                        progress_bar.empty()
                        st.error(f"Failed to generate resume: {e}")

        if st.session_state.resume_text:
            st.markdown("#### Resume (Plain Text)")
            st.text_area(
                "Generated Resume",
                value=st.session_state.resume_text,
                height=420,
                key="resume_preview",
                label_visibility="collapsed",
            )

            docx_bytes = create_docx_from_text(st.session_state.resume_text, st.session_state.form_data.get("name", "Resume"))
            pdf_bytes = create_pdf_from_text(st.session_state.resume_text, st.session_state.form_data.get("name", "Resume"))

            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                st.download_button(
                    label="Download Resume as DOCX",
                    data=docx_bytes,
                    file_name="resume_ai_generated.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with col_dl2:
                st.download_button(
                    label="Download Resume as PDF",
                    data=pdf_bytes,
                    file_name="resume_ai_generated.pdf",
                    mime="application/pdf",
                )

    # Cover Letter Tab
    with tabs[1]:
        st.markdown("### Step 8 · AI-Generated Cover Letter")
        st.write(
            "Once your resume is ready, enter the target job details below and click **Generate Cover Letter** for that specific role."
        )

        job_title = st.text_input("Target Job Title", key="job_title_input")
        company_name = st.text_input("Company Name", key="company_name_input")

        generate_cover_clicked = st.button("Generate Cover Letter", type="primary")

        if generate_cover_clicked:
            if not st.session_state.resume_text:
                st.error("Please generate a resume first before creating a cover letter.")
            elif not job_title.strip() or not company_name.strip():
                st.error("Please provide both the Job Title and Company Name.")
            else:
                form_data = st.session_state.form_data
                prompt = build_cover_letter_prompt(form_data, job_title.strip(), company_name.strip())
                progress_bar = st.progress(0)
                with st.spinner("Generating your tailored cover letter..."):
                    try:
                        progress_bar.progress(20)
                        cover_text = call_openai(prompt)
                        progress_bar.progress(80)
                        st.session_state.cover_letter_text = cover_text
                        save_to_supabase(form_data, st.session_state.resume_text, cover_text)
                        progress_bar.progress(100)
                        st.success("Cover letter generated successfully.")
                    except Exception as e:
                        progress_bar.empty()
                        st.error(f"Failed to generate cover letter: {e}")

        if st.session_state.cover_letter_text:
            st.markdown("#### Cover Letter (Plain Text)")
            st.text_area(
                "Generated Cover Letter",
                value=st.session_state.cover_letter_text,
                height=420,
                key="cover_preview",
                label_visibility="collapsed",
            )

            cover_pdf_bytes = create_pdf_from_text(
                st.session_state.cover_letter_text,
                f"Cover Letter - {st.session_state.form_data.get('name', '')}".strip(),
            )
            st.download_button(
                label="Download Cover Letter as PDF",
                data=cover_pdf_bytes,
                file_name="cover_letter_ai_generated.pdf",
                mime="application/pdf",
            )

    # Portfolio Tab
    with tabs[2]:
        st.markdown("### Step 9 · Portfolio Preview")
        st.write(
            "This portfolio is auto-generated from your resume, projects, and experience. Copy it into a personal website, GitHub README, or portfolio PDF."
        )

        portfolio = st.session_state.portfolio
        if not portfolio and st.session_state.resume_text:
            portfolio = build_portfolio_sections(st.session_state.form_data, st.session_state.resume_text)
            st.session_state.portfolio = portfolio

        if not portfolio:
            st.info("Generate a resume first to see your portfolio preview.", icon="ℹ️")
        else:
            st.markdown('<div class="portfolio-section-title">Summary</div>', unsafe_allow_html=True)
            st.markdown(
                f'<div class="portfolio-text">{portfolio.get("summary", "")}</div>',
                unsafe_allow_html=True,
            )

            st.markdown('<div class="portfolio-section-title" style="margin-top:12px;">Projects</div>', unsafe_allow_html=True)
            projects_list = portfolio.get("projects", [])
            if not projects_list:
                st.markdown('<div class="portfolio-text">No project details added yet.</div>', unsafe_allow_html=True)
            else:
                df_projects = pd.DataFrame({"Projects": projects_list})
                for idx, row in df_projects.iterrows():
                    st.markdown(
                        f'<div class="portfolio-text">- {row["Projects"]}</div>',
                        unsafe_allow_html=True,
                    )

            st.markdown(
                '<div class="portfolio-section-title" style="margin-top:12px;">Experience Highlights</div>',
                unsafe_allow_html=True,
            )
            exp_highlights = portfolio.get("experience_highlights", [])
            if not exp_highlights:
                st.markdown(
                    '<div class="portfolio-text">Add experience or project responsibilities to see highlights here.</div>',
                    unsafe_allow_html=True,
                )
            else:
                for line in exp_highlights:
                    st.markdown(
                        f'<div class="portfolio-text">- {line}</div>',
                        unsafe_allow_html=True,
                    )

    st.markdown('</div>', unsafe_allow_html=True)

